from __future__ import annotations

from io import BytesIO
from pathlib import Path
from typing import Any, Callable, Dict, List, Optional, Sequence

from django.conf import settings

from docx import Document
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Mm, Pt, RGBColor

from .utils import fmt_date, fmt_time
from . import report_config as cfg


DEFAULT_TEMPLATE_PATH = Path(getattr(settings, "BASE_DIR", Path("."))) / "api" / "assets" / "Modelo.docx"
EMU_PER_MM = 36000  # 1mm = 36000 EMU


# ---------------------------------------------------------------------------
# Helpers internos de baixo nível
# ---------------------------------------------------------------------------

def _hex_to_rgb(hex_color: str) -> RGBColor:
    hx = (hex_color or "").strip().lstrip("#").upper()
    if len(hx) != 6:
        hx = "000000"
    return RGBColor.from_string(hx)


def _set_cell_shading(cell, fill_hex: str) -> None:
    fill = (fill_hex or "").strip().lstrip("#")
    if not fill:
        return
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def _set_cell_text(
    cell,
    text: str,
    *,
    bold: bool = False,
    underline: bool = False,
    color_hex: Optional[str] = None,
    align: Optional[int] = None,
    font_size: Pt = Pt(9),
) -> None:
    cell.text = ""
    p = cell.paragraphs[0]
    if align is not None:
        p.alignment = align
    run = p.add_run("" if text is None else str(text))
    run.bold = bold
    run.underline = underline
    run.font.size = font_size
    if color_hex:
        run.font.color.rgb = _hex_to_rgb(color_hex)


def _add_field(paragraph, instr: str, *, font_size: Pt = Pt(8)) -> None:
    fld = OxmlElement("w:fldSimple")
    fld.set(qn("w:instr"), instr)

    r = OxmlElement("w:r")
    r_pr = OxmlElement("w:rPr")

    sz = OxmlElement("w:sz")
    sz.set(qn("w:val"), str(int(font_size.pt * 2)))  # half-points
    r_pr.append(sz)
    r.append(r_pr)

    t = OxmlElement("w:t")
    t.text = "1"  # placeholder; Word substitui ao atualizar campos
    r.append(t)

    fld.append(r)
    paragraph._p.append(fld)


# ---------------------------------------------------------------------------
# Infraestrutura de documento
# ---------------------------------------------------------------------------

def _apply_page_number_footer(doc: Document) -> None:
    try:
        section = doc.sections[0]
    except Exception:
        return

    footer = section.footer
    p = footer.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT

    r1 = p.add_run("Página ")
    r1.font.size = Pt(8)

    _add_field(p, "PAGE", font_size=Pt(8))

    r2 = p.add_run(" de ")
    r2.font.size = Pt(8)

    _add_field(p, "NUMPAGES", font_size=Pt(8))


def _clear_document_body(doc: Document) -> None:
    body = doc._element.body
    for child in list(body.iterchildren()):
        if child.tag == qn("w:sectPr"):
            continue
        body.remove(child)


def _usable_width_mm(doc: Document) -> float:
    section = doc.sections[0]
    return float(section.page_width - section.left_margin - section.right_margin) / EMU_PER_MM


def _calc_col_widths_mm(doc: Document, weights: List[float]) -> List[float]:
    wmm = _usable_width_mm(doc)
    total = sum(weights) if weights else 1
    return [(wmm * (w / total)) for w in weights]


def _add_table(doc: Document, *, rows: int, cols: int, col_widths_mm: List[float]):
    table = doc.add_table(rows=rows, cols=cols)
    try:
        table.style = "Table Grid"
    except Exception:
        pass

    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False

    for i, w in enumerate(col_widths_mm):
        for cell in table.columns[i].cells:
            cell.width = Mm(w)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER

    return table


def _render_header_row(table, headers: List[str], *, fill_hex: str = cfg.HEADER_FILL, font_size: Pt = Pt(9)) -> None:
    for j, h in enumerate(headers):
        cell = table.cell(0, j)
        _set_cell_shading(cell, fill_hex)
        _set_cell_text(cell, h, bold=True, font_size=font_size)


def _add_total_row(doc: Document, total: int) -> None:
    widths = _calc_col_widths_mm(doc, [80, 20])
    tbl = _add_table(doc, rows=1, cols=2, col_widths_mm=widths)

    for j in range(2):
        _set_cell_shading(tbl.cell(0, j), cfg.HEADER_FILL)

    _set_cell_text(tbl.cell(0, 0), "Total", bold=True, underline=True, font_size=Pt(10))
    _set_cell_text(
        tbl.cell(0, 1),
        str(total),
        bold=True,
        underline=True,
        align=WD_ALIGN_PARAGRAPH.RIGHT,
        font_size=Pt(10),
    )


def _init_doc(title: str, template_path: Optional[Path] = None) -> Document:
    template_path = template_path or DEFAULT_TEMPLATE_PATH
    doc = Document(str(template_path)) if template_path and Path(template_path).exists() else Document()

    _clear_document_body(doc)
    _apply_page_number_footer(doc)

    heading = doc.add_paragraph(title)
    try:
        heading.style = "Heading 2"
    except Exception:
        pass

    return doc


def _finalize_doc(doc: Document, total: int) -> bytes:
    doc.add_paragraph("")
    _add_total_row(doc, total=total)

    buf = BytesIO()
    doc.save(buf)
    return buf.getvalue()


# ---------------------------------------------------------------------------
# Relatório flat genérico (tabela simples com header + N linhas de dados)
# ---------------------------------------------------------------------------

def _build_flat_docx(
    title: str,
    rows: Sequence[Dict[str, Any]],
    headers: List[str],
    weights: List[int],
    row_builder: Callable,
    template_path: Optional[Path],
    *,
    font_size: Pt = Pt(8),
) -> bytes:
    """Flat report: row_builder(tbl, row_index, row_data) writes cells directly."""
    doc = _init_doc(title, template_path)

    if not rows:
        doc.add_paragraph("Nenhum registro encontrado para os filtros aplicados.")
        buf = BytesIO()
        doc.save(buf)
        return buf.getvalue()

    colw = _calc_col_widths_mm(doc, weights)
    tbl = _add_table(doc, rows=1 + len(rows), cols=len(headers), col_widths_mm=colw)

    _render_header_row(tbl, headers, font_size=font_size)

    for i, r in enumerate(rows, start=1):
        row_builder(tbl, i, r)

    return _finalize_doc(doc, total=len(rows))


# ---------------------------------------------------------------------------
# Relatórios públicos — flat
# ---------------------------------------------------------------------------

def gerar_relatorio_operadores(
    operadores: Sequence[Dict[str, Any]],
    template_path: Optional[Path] = None,
) -> bytes:
    def row_builder(tbl, i, op):
        nome = op.get("nome_completo") or op.get("nome") or "--"
        email = op.get("email") or "--"
        _set_cell_text(tbl.cell(i, 0), str(nome), align=WD_ALIGN_PARAGRAPH.LEFT, font_size=Pt(9))
        _set_cell_text(tbl.cell(i, 1), str(email), align=WD_ALIGN_PARAGRAPH.LEFT, font_size=Pt(9))

    return _build_flat_docx(
        "Operadores de Áudio", operadores,
        headers=["Nome", "E-mail"],
        weights=cfg.COLS_OPERADORES,
        row_builder=row_builder,
        template_path=template_path,
        font_size=Pt(9),
    )


def gerar_relatorio_operacoes_entradas(
    entradas: Sequence[Dict[str, Any]],
    template_path: Optional[Path] = None,
) -> bytes:
    def row_builder(tbl, i, r):
        anom = bool(r.get("anormalidade"))
        anom_txt = "SIM" if anom else "Não"
        anom_color = cfg.COLOR_RED if anom else cfg.COLOR_GREEN

        cells = [
            (r.get("sala") or "--", WD_ALIGN_PARAGRAPH.LEFT, True, None),
            (fmt_date(r.get("data")), WD_ALIGN_PARAGRAPH.LEFT, False, None),
            (r.get("operador") or "--", WD_ALIGN_PARAGRAPH.LEFT, False, None),
            (r.get("tipo") or "--", WD_ALIGN_PARAGRAPH.LEFT, False, None),
            (r.get("evento") or "--", WD_ALIGN_PARAGRAPH.LEFT, False, None),
            (fmt_time(r.get("pauta")), WD_ALIGN_PARAGRAPH.CENTER, False, None),
            (fmt_time(r.get("inicio")), WD_ALIGN_PARAGRAPH.CENTER, False, None),
            (fmt_time(r.get("fim")), WD_ALIGN_PARAGRAPH.CENTER, False, None),
            (anom_txt, WD_ALIGN_PARAGRAPH.CENTER, True, anom_color),
        ]

        for j, (txt, align, bold, color) in enumerate(cells):
            _set_cell_text(tbl.cell(i, j), str(txt), bold=(bold if j in (0, 8) else False), color_hex=color, align=align, font_size=Pt(8))

    return _build_flat_docx(
        "Registros de Operação (Entradas)", entradas,
        headers=["Local", "Data", "Operador", "Tipo", "Evento", "Pauta", "Início", "Fim", "Anormalidade?"],
        weights=cfg.COLS_OPERACOES_ENTRADAS,
        row_builder=row_builder,
        template_path=template_path,
    )


def gerar_relatorio_anormalidades(
    anormalidades: Sequence[Dict[str, Any]],
    template_path: Optional[Path] = None,
) -> bytes:
    def row_builder(tbl, i, r):
        solucionada = bool(r.get("solucionada"))
        preju = bool(r.get("houve_prejuizo"))
        recl = bool(r.get("houve_reclamacao"))

        sol_txt = "Sim" if solucionada else "Não"
        sol_color = cfg.COLOR_GREEN if solucionada else cfg.COLOR_RED

        prej_txt = "Sim" if preju else "Não"
        prej_color = cfg.COLOR_RED if preju else cfg.COLOR_MUTED

        recl_txt = "Sim" if recl else "Não"
        recl_color = cfg.COLOR_RED if recl else cfg.COLOR_MUTED

        cells = [
            (fmt_date(r.get("data")), None),
            (r.get("sala") or "--", None),
            (r.get("registrado_por") or "--", None),
            (r.get("descricao") or "--", None),
            (sol_txt, sol_color),
            (prej_txt, prej_color),
            (recl_txt, recl_color),
        ]

        for j, (txt, color) in enumerate(cells):
            align = WD_ALIGN_PARAGRAPH.CENTER if j >= 4 else WD_ALIGN_PARAGRAPH.LEFT
            bold = j >= 4
            _set_cell_text(tbl.cell(i, j), str(txt), bold=bold, color_hex=color, align=align, font_size=Pt(8))

    return _build_flat_docx(
        "Relatórios de Anormalidades", anormalidades,
        headers=["Data", "Local", "Registrado por", "Descrição", "Solucionada", "Prejuízo", "Reclamação"],
        weights=cfg.COLS_ANORMALIDADES,
        row_builder=row_builder,
        template_path=template_path,
    )


# ---------------------------------------------------------------------------
# Relatórios públicos — master/detail
# ---------------------------------------------------------------------------

def gerar_relatorio_operacoes_sessoes(
    sessoes: Sequence[Dict[str, Any]],
    template_path: Optional[Path] = None,
) -> bytes:
    doc = _init_doc("Registros de Operação (Sessões)", template_path)

    if not sessoes:
        doc.add_paragraph("Nenhum registro encontrado para os filtros aplicados.")
        buf = BytesIO()
        doc.save(buf)
        return buf.getvalue()

    master_cols = _calc_col_widths_mm(doc, cfg.COLS_OPERACOES_SESSOES_MASTER)
    ent_cols = _calc_col_widths_mm(doc, cfg.COLS_OPERACOES_SESSOES_ENTRADAS)

    for idx, s in enumerate(sessoes):
        t_master = _add_table(doc, rows=2, cols=5, col_widths_mm=master_cols)
        _render_header_row(t_master, ["Local", "Data", "1º Registro por", "Checklist?", "Em Aberto?"])

        sala_txt = s.get("sala") or "--"
        data_txt = fmt_date(s.get("data"))
        autor_txt = s.get("autor") or "--"

        verific_txt = str(s.get("verificacao") or "--").strip() or "--"
        verific_color = cfg.COLOR_GREEN if verific_txt.lower() == "realizado" else cfg.COLOR_MUTED

        em_txt = str(s.get("em_aberto") or "--").strip() or "--"
        em_color = cfg.COLOR_BLUE if em_txt.lower() == "sim" else cfg.COLOR_DARK

        values = [
            (str(sala_txt), True, None, WD_ALIGN_PARAGRAPH.LEFT),
            (str(data_txt), True, None, WD_ALIGN_PARAGRAPH.LEFT),
            (str(autor_txt), True, None, WD_ALIGN_PARAGRAPH.LEFT),
            (str(verific_txt), True, verific_color, WD_ALIGN_PARAGRAPH.LEFT),
            (str(em_txt), True, em_color, WD_ALIGN_PARAGRAPH.LEFT),
        ]

        for j, (txt, bold, color, align) in enumerate(values):
            cell = t_master.cell(1, j)
            _set_cell_shading(cell, cfg.DATA_ROW_FILL)
            _set_cell_text(cell, txt, bold=bold, color_hex=color, align=align, font_size=Pt(9))

        doc.add_paragraph("")

        bar = _add_table(doc, rows=1, cols=1, col_widths_mm=_calc_col_widths_mm(doc, [100]))
        _set_cell_shading(bar.cell(0, 0), cfg.DETAIL_BAR_FILL)
        _set_cell_text(bar.cell(0, 0), "Entradas da Operação:", bold=True, font_size=Pt(9))

        entradas = s.get("entradas") or []
        t_ent = _add_table(doc, rows=1 + (len(entradas) if entradas else 1), cols=8, col_widths_mm=ent_cols)
        _render_header_row(t_ent, ["Nº", "Operador", "Tipo", "Evento", "Pauta", "Início", "Fim", "Anormalidade?"],
                           fill_hex=cfg.HEADER_DETAIL_FILL, font_size=Pt(8))

        if entradas:
            for i, ent in enumerate(entradas, start=1):
                ordem = ent.get("ordem")
                ordem_txt = f"{ordem}º" if ordem is not None and str(ordem) != "" else "--"

                anom = bool(ent.get("anormalidade"))
                anom_txt = "SIM" if anom else "Não"
                anom_color = cfg.COLOR_RED if anom else cfg.COLOR_GREEN

                row_vals = [
                    (ordem_txt, None, WD_ALIGN_PARAGRAPH.CENTER),
                    (ent.get("operador") or "--", None, WD_ALIGN_PARAGRAPH.LEFT),
                    (ent.get("tipo") or "--", None, WD_ALIGN_PARAGRAPH.LEFT),
                    (ent.get("evento") or "--", None, WD_ALIGN_PARAGRAPH.LEFT),
                    (fmt_time(ent.get("pauta")), None, WD_ALIGN_PARAGRAPH.CENTER),
                    (fmt_time(ent.get("inicio")), None, WD_ALIGN_PARAGRAPH.CENTER),
                    (fmt_time(ent.get("fim")), None, WD_ALIGN_PARAGRAPH.CENTER),
                    (anom_txt, anom_color, WD_ALIGN_PARAGRAPH.CENTER),
                ]

                for j, (txt, color, align) in enumerate(row_vals):
                    _set_cell_text(
                        t_ent.cell(i, j),
                        str(txt),
                        color_hex=color,
                        align=align,
                        font_size=Pt(8),
                        bold=(j == 7),
                    )
        else:
            c = t_ent.cell(1, 0)
            _set_cell_text(c, "Nenhuma entrada registrada nesta sessão.", font_size=Pt(8))
            for j in range(1, 8):
                c = c.merge(t_ent.cell(1, j))

        if idx < len(sessoes) - 1:
            doc.add_paragraph("")
            doc.add_paragraph("")

    return _finalize_doc(doc, total=len(sessoes))


def gerar_relatorio_checklists(
    checklists: Sequence[Dict[str, Any]],
    template_path: Optional[Path] = None,
) -> bytes:
    doc = _init_doc("Verificação de Plenários", template_path)

    if not checklists:
        doc.add_paragraph("Nenhum registro encontrado para os filtros aplicados.")
        buf = BytesIO()
        doc.save(buf)
        return buf.getvalue()

    master_cols = _calc_col_widths_mm(doc, cfg.COLS_CHECKLISTS_MASTER)
    itens_cols = _calc_col_widths_mm(doc, cfg.COLS_CHECKLISTS_ITENS)

    for idx, chk in enumerate(checklists):
        itens = chk.get("itens") or []

        has_failure = any(str(it.get("status") or "").strip().lower() == "falha" for it in itens)
        status_txt = "Falha" if has_failure else "Ok"
        status_color = cfg.COLOR_RED if has_failure else cfg.COLOR_GREEN

        t_master = _add_table(doc, rows=2, cols=7, col_widths_mm=master_cols)
        _render_header_row(t_master, ["Local", "Data", "Verificado por", "Início", "Término", "Duração", "Status"], font_size=Pt(8))

        sala = chk.get("sala_nome") or chk.get("sala") or "--"
        data_txt = fmt_date(chk.get("data"))
        operador = chk.get("operador") or "--"
        inicio = fmt_time(chk.get("inicio"))
        termino = fmt_time(chk.get("termino"))
        duracao = chk.get("duracao") or "--"

        row_vals = [
            (str(sala), True, None, WD_ALIGN_PARAGRAPH.LEFT),
            (str(data_txt), False, None, WD_ALIGN_PARAGRAPH.LEFT),
            (str(operador), False, None, WD_ALIGN_PARAGRAPH.LEFT),
            (str(inicio), False, None, WD_ALIGN_PARAGRAPH.CENTER),
            (str(termino), False, None, WD_ALIGN_PARAGRAPH.CENTER),
            (str(duracao), False, None, WD_ALIGN_PARAGRAPH.CENTER),
            (str(status_txt), True, status_color, WD_ALIGN_PARAGRAPH.CENTER),
        ]

        for j, (txt, bold, color, align) in enumerate(row_vals):
            cell = t_master.cell(1, j)
            _set_cell_shading(cell, cfg.DATA_ROW_FILL)
            _set_cell_text(cell, txt, bold=bold, color_hex=color, align=align, font_size=Pt(8))

        doc.add_paragraph("")

        bar = _add_table(doc, rows=1, cols=1, col_widths_mm=_calc_col_widths_mm(doc, [100]))
        _set_cell_shading(bar.cell(0, 0), cfg.DETAIL_BAR_FILL)
        _set_cell_text(bar.cell(0, 0), "Detalhes da Verificação:", bold=True, font_size=Pt(9))

        rows_count = 1 + (len(itens) if itens else 1)
        t_itens = _add_table(doc, rows=rows_count, cols=3, col_widths_mm=itens_cols)
        _render_header_row(t_itens, ["Item verificado", "Status", "Descrição"],
                           fill_hex=cfg.HEADER_DETAIL_FILL, font_size=Pt(8))

        if itens:
            for i, it in enumerate(itens, start=1):
                is_text = (it.get("tipo_widget") or "radio") == "text"

                if is_text:
                    st = "Texto"
                    st_color = cfg.COLOR_SLATE
                    desc_txt = it.get("valor_texto")
                    desc_txt = desc_txt if (desc_txt is not None and str(desc_txt).strip() != "") else "-"
                else:
                    st = str(it.get("status") or "--").strip() or "--"
                    st_lower = st.lower()
                    if st_lower == "falha":
                        st_color = cfg.COLOR_RED
                    elif st_lower == "ok":
                        st_color = cfg.COLOR_GREEN
                    else:
                        st_color = cfg.COLOR_SLATE
                    desc_txt = it.get("falha")
                    desc_txt = desc_txt if (desc_txt is not None and str(desc_txt).strip() != "") else "-"

                _set_cell_text(t_itens.cell(i, 0), str(it.get("item") or "--"),
                               align=WD_ALIGN_PARAGRAPH.LEFT, font_size=Pt(8))
                _set_cell_text(t_itens.cell(i, 1), st, bold=True, color_hex=st_color,
                               align=WD_ALIGN_PARAGRAPH.CENTER, font_size=Pt(8))
                _set_cell_text(t_itens.cell(i, 2), str(desc_txt),
                               align=WD_ALIGN_PARAGRAPH.LEFT, font_size=Pt(8))
        else:
            _set_cell_text(t_itens.cell(1, 0), "Nenhum item encontrado.", font_size=Pt(8))
            t_itens.cell(1, 0).merge(t_itens.cell(1, 1))
            t_itens.cell(1, 0).merge(t_itens.cell(1, 2))

        if idx < len(checklists) - 1:
            doc.add_paragraph("")
            doc.add_paragraph("")

    return _finalize_doc(doc, total=len(checklists))
